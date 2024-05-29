from django.shortcuts import render
from rest_framework.response import Response
from rest_framework.decorators import api_view,permission_classes
from .serializer import *
from rest_framework import status
from rest_framework_simplejwt.tokens import RefreshToken

from rest_framework.views import APIView
from django.contrib.auth import authenticate
from django.http import JsonResponse
from .models import *
# Create your views here.
@api_view(["GET"])
def getRoutes(request):
    return Response("Hello world")


@api_view(["POST"])
def RegisterUser(request):
    user_data = request.data.copy()  # Make a copy of the data to add 'role'
    user_data.pop("role")
    user_data.pop("teams")
    user_serializer= UserSerializer(data=user_data)
    if user_serializer.is_valid():
        user=User.objects.create_user(
            first_name=user_data["first_name"],
            username=user_data["username"],
            last_name=user_data["last_name"],
            email=user_data["email"],
            password=user_data["password"],
        )
        if request.data["role"]=="team lead":
            team = Team.objects.create(lead=user)
            role=UserRole.objects.create(user=user,role=request.data["role"],teams=team)
        elif request.data["role"]=="developer":
            team = Team.objects.get(id=request.data["teams"])
            role=UserRole.objects.create(user=user,role=request.data["role"],teams=team)
        else:
            role=UserRole.objects.create(user=user,role=request.data["role"])
        role_serializer=RoleSerializer(role)
        return Response(role_serializer.data,status=status.HTTP_201_CREATED)
    return Response(user_serializer.errors,status=status.HTTP_400_BAD_REQUEST)

@api_view(["GET"])
def GetUsers(request):
    user=User.objects.all()
    user_serializer=UserSerializer(user,many=True)
    return Response(user_serializer.data,status=status.HTTP_200_OK)

@api_view(["GET"])
def getUserRole(request):
    user=UserRole.objects.all()
    user_serializer=RoleSerializer(user,many=True)
    return Response(user_serializer.data,status=status.HTTP_200_OK)


class Loginview(APIView):
    def post(self,request):
        username=request.data['username']
        password=request.data['password']
        user = authenticate(username=username,password=password)
        if user is None:
            return (Response({"error":"Invalid Credentials"},status=status.HTTP_400_BAD_REQUEST))
        refresh=RefreshToken.for_user(user)
        role=UserRole.objects.get(user=user)
        user=RoleSerializer(role,many=False).data
        user["refresh_token"]=str(refresh)
        user["access_token"]=str(refresh.access_token)
        return JsonResponse(user,status=status.HTTP_200_OK)
    

@api_view(["GET"])
def GetProjects(request):
    project=Project.objects.all()
    project_serializer=ProjectSerializer(project,many=True)
    return Response(project_serializer.data,status=status.HTTP_200_OK)

@api_view(["POST","PUT"])
def CreateProject(request):
    if request.method=="POST":
        try:
            data = request.data.copy()
            try:
                project_serializer=ProjectSerializer(data=data)
            except Exception as e:
                print(e)
            if project_serializer.is_valid():
                project_serializer.save()
                return Response(project_serializer.data,status=status.HTTP_200_OK)
            return Response(project_serializer.errors,status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            print(e)
    elif request.method=="PUT":
        print("PUT---")
        try:
            print("REACHED PUT")
            print(request.data["id"])
            project_id = request.data.pop('id') 
            project = Project.objects.get(pk=project_id)
            data = request.data.copy()
            project_serializer = ProjectSerializer(project, data=data, partial=True)
            if project_serializer.is_valid():
                project_serializer.save()
                return Response(project_serializer.data, status=status.HTTP_200_OK)
            return Response(project_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except (KeyError, Project.DoesNotExist):
            return Response({"error": "Project does not exist"}, status=status.HTTP_404_NOT_FOUND)
    return Response({"details":"error"},status=status.HTTP_400_BAD_REQUEST)
        
        

@api_view(["PUT"])
def UpdateProjectStatus(request):
    try:
        print(request.data["id"])
        project_id = int(request.data.pop('id') )
        project = Project.objects.get(pk=project_id)
        project.column=request.data["status"]
        #done doing todo
        project.save()
        return Response({"success": "updated"}, status=status.HTTP_202_ACCEPTED)
    except (KeyError, Project.DoesNotExist):
        return Response({"error": "Project does not exist"}, status=status.HTTP_404_NOT_FOUND)
    

@api_view(["GET","POST","PUT"])
def TaskView(request):
    if request.method=="GET":
        task=Task.objects.all()
        task_serializer=TaskSerializer(task,many=True)
        return Response(task_serializer.data,status=status.HTTP_200_OK)
    elif request.method == "POST":
        print("POST HITTED TASK VIEW")
        print(request.data)
        data = request.data.copy()
        task_serializer=TaskSerializer(data=data)
        project=Project.objects.get(id=request.data["pid"])
        project.assigned=True
        project.save()
        if task_serializer.is_valid():
            task_serializer.save()
            return Response(task_serializer.data,status=status.HTTP_201_CREATED)
        return Response(task_serializer.errors,status=status.HTTP_400_BAD_REQUEST)
    elif request.method=="PUT":
        print("PUT---")
        try:
            print("REACHED PUT")
            print(request.data["id"])
            task_id = request.data.pop('id') 
            task = Task.objects.get(pk=task_id)
            data = request.data.copy()
            task_serializer = TaskSerializer(task, data=data, partial=True)
            if task_serializer.is_valid():
                task_serializer.save()
                return Response(task_serializer.data, status=status.HTTP_200_OK)
            return Response(task_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except (KeyError, Task.DoesNotExist):
            return Response({"error": "Task does not exist"}, status=status.HTTP_404_NOT_FOUND)
    return Response({"details":"error"},status=status.HTTP_400_BAD_REQUEST)


@api_view(["GET"])
def TaskbyId(request,id):
    # id = int(request.data["id"])
    data=Task.objects.filter(assigned_to=id)
    task=TaskSerializer(data,many=True)
    try:
        return Response(task.data,status=status.HTTP_200_OK)
    except Exception as e:
        return Response(e,status=status.HTTP_400_BAD_REQUEST)

@api_view(["PUT"])
def UpdateTaskStatus(request):
    try:
        print(request.data["id"])
        task_id = int(request.data.pop('id') )
        task = Task.objects.get(pk=task_id)
        task.status=request.data["status"]
        #done doing todo
        project=Project.objects.get(pk=task.pid.id)
        project.column=request.data["status"]
        project.save()
        task.save()
        return Response({"success": "updated"}, status=status.HTTP_202_ACCEPTED)
    except (KeyError, Task.DoesNotExist):
        return Response({"error": "Project does not exist"}, status=status.HTTP_404_NOT_FOUND)
    

@api_view(["GET"])
def TaskCompleted(request):
    task=Task.objects.filter(status="done")
    task_serializer=TaskSerializer(task,many=True)
    return Response(task_serializer.data,status=status.HTTP_200_OK)

@api_view(["GET"])
def GetNaProjects(request):
    project=Project.objects.filter(assigned=False)
    project_serializer=ProjectSerializer(project,many=True)
    return Response(project_serializer.data,status=status.HTTP_200_OK)

@api_view(["GET"])
def getDev(request):
    user=UserRole.objects.filter(role="team lead")
    user_serializer=RoleSerializer(user,many=True)
    return Response(user_serializer.data,status=status.HTTP_200_OK)


@api_view(["GET"])
def GetSaleProject(request,id):
    project=Project.objects.filter(bid_by=id)
    project_serializer=ProjectSerializer(project,many=True)
    return Response(project_serializer.data,status=status.HTTP_200_OK)


@api_view(["GET","POST","PUT"])
def getUserStory(request):
    if request.method=="GET":
        userStory=UserStory.objects.all()
        story_serializer=UserStorySerializer(userStory,many=True)
        return Response(story_serializer.data,status=status.HTTP_200_OK)
    elif request.method=="POST":
        ustory=UserStorySerializer(data=request.data)
        if ustory.is_valid():
            ustory.save()
            return Response(ustory.data, status=status.HTTP_200_OK)
        return Response(ustory.errors, status=status.HTTP_400_BAD_REQUEST)
    return Response({"details":"error"},status=status.HTTP_400_BAD_REQUEST)
    


@api_view(["GET"])
def getUserStorybyId(request,id):
    userstory=UserStory.objects.filter(pid=id)
    story_serializer=UserStorySerializer(userstory,many=True)
    return Response(story_serializer.data,status=status.HTTP_200_OK)

# @api_view(["GET"])
# def 







# myapp/views.py
from django.http import JsonResponse, HttpResponse
from django.core.files.base import ContentFile
from langchain.prompts import PromptTemplate
from langchain_google_genai import GoogleGenerativeAI
from docx import Document
from io import BytesIO
from django.views.decorators.csrf import csrf_exempt
from moviepy.editor import VideoFileClip
import os
from pydub import AudioSegment
from django.core.files.storage import default_storage

llm = GoogleGenerativeAI(model="gemini-pro", google_api_key="AIzaSyCTJGEfrQMEZ1S9DKHl9_Hr_pdx3_7gAi4")

template = """Question: {question}

Answer: Let's think step by step and generate requirements and Generate some user story from this and shape the document according to agile developement  
on this tempelate Software Requirements Specification (SRS) Template
please be detailed in each section for more clearification and please be detailed in defining installation requirements like what can be
required to install in this case of an application if text is urdu translate it into english then generate SRS
1. Introduction

1.1 System Name
1.2 Version
1.3 Author(s)
1.4 Date
1.5 Purpose
1.6 Scope
1.7 Definitions, Acronyms, and Abbreviations

2. Overall Description

2.1 Product Overview
2.2 Product Functions
2.3 User Characteristics
2.4 Operating Environment
2.5 Design and Implementation Constraints
2.6 Assumptions and Dependencies

3. System Requirements

3.1 Functional Requirements
3.1.1 [Module Name]
3.1.2 [Module Name]
3.1.3 ...
3.2 Non-Functional Requirements
3.2.1 Performance Requirements
3.2.2 Security Requirements
3.2.3 Reliability Requirements
3.2.4 Maintainability Requirements

4. External Interface Requirements

4.1 User Interfaces
4.2 Hardware Interfaces
4.3 Software Interfaces
4.4 Communications Interfaces

5. System Features

5.1 Feature List
5.2 Feature Descriptions
5.3 Priority and Dependencies

6. Other Requirements

6.1 Data Requirements
6.2 Packaging and Distribution Requirements
6.3 Installation Requirements
6.4 Support Requirements
"""

template2 = """Question: {question}

Answer: Generate user stories atleast 8 more from the following SRS don't put star in it and seperate them by lines only

"""

prompt = PromptTemplate.from_template(template)
prompt2=PromptTemplate.from_template(template2)

chain = prompt | llm
chain2=prompt2 | llm

def create_word_document(text):
    document = Document()
    document.add_paragraph(text)
    return document

def Break(start,end,text):
    x=text.find(start)
    y=text.find(end)
    title=text[x:y]
    title=title.split("**")
    return(title[-1])

@api_view(["POST"])
def generate_requirements(request):
    if request.method == 'POST':
        word_file = request.FILES['wordFile']
        document = Document(BytesIO(word_file.read()))
        text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
        requirements = chain.invoke({"question": text})
        title=Break("1.1 System Name","**1.2",requirements)
        description=Break("1.6 Scope","**1.7",requirements)
        user=User.objects.get(id=request.data["id"])
        pid=Project.objects.create(
            title=title.strip(),
            description=description.strip(),
            bid_by=user,
            column="todo"
        )
        print(title.strip(),description.strip(),"HELLo")
        story = chain2.invoke({"question": requirements})
        story=story.split("\n")
        for i in story:
            UserStory.objects.create(
                description=i[1:].strip(),
                pid=pid,
            )
            print("HELOO THIS IS RECORD",i[1:])
        doc = create_word_document(requirements)
        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        response = HttpResponse(output_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=SRS.docx'
        return response
    return HttpResponse(status=405)


def check(request):
    if request.method == 'POST':
        word_file = request.FILES['wordFile']
        document = Document(BytesIO(word_file.read()))
        text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
        print(text)
        return JsonResponse({"message": "success"})
    return HttpResponse(status=405)

def convert_mp4_to_wav(request):
    if request.method == 'POST':
        try:
            if 'file' not in request.FILES:
                return HttpResponse('No file part', status=400)

            file = request.FILES['file']
            if file.name == '' or not file.name.endswith('.mp4'):
                return HttpResponse('Invalid file format. Please upload an MP4 file.', status=400)

            video_clip = VideoFileClip(file)
            wav_filepath = 'output.wav'
            video_clip.audio.write_audiofile(wav_filepath)
            video_clip.close()
            response = HttpResponse(open(wav_filepath, 'rb'), content_type='audio/wav')
            response['Content-Disposition'] = 'attachment; filename=output.wav'
            os.remove(wav_filepath)
            return response
        except Exception as e:
            print("except")
            return HttpResponse(str(e), status=500)
    return HttpResponse(status=405)


def append_wav_files(request):
    if request.method == 'POST':
        try:
            if 'file1' not in request.FILES or 'file2' not in request.FILES:
                return HttpResponse('Two files are required', status=400)

            file1 = request.FILES['file1']
            file2 = request.FILES['file2']
            if not file1.name.endswith('.wav') or not file2.name.endswith('.wav'):
                return HttpResponse('Invalid file format. Please upload two WAV files.', status=400)

            audio1 = AudioSegment.from_wav(file1)
            audio2 = AudioSegment.from_wav(file2)
            combined_audio = audio1 + audio2

            combined_filepath = 'combined_audio.wav'
            combined_audio.export(combined_filepath, format="wav")
            response = HttpResponse(open(combined_filepath, 'rb'), content_type='audio/wav')
            response['Content-Disposition'] = 'attachment; filename=combined_audio.wav'
            os.remove(combined_filepath)
            return response
        except Exception as e:
            return HttpResponse(str(e), status=500)
    return HttpResponse(status=405)
@api_view(['POST'])
def create_pteam(request):
    if request.method == 'POST':
        serializer = PTeamSerializer(data=request.data)
        if serializer.is_valid():
            project=Project.objects.get(id=request.data["pid"])
            project.assigned=True
            project.save()
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
@api_view(['GET'])
def list_pteams(request):
    if request.method == 'GET':
        pteams = PTeam.objects.all()
        serializer = PTeamSerializer(pteams, many=True)
        return Response(serializer.data)
    
from django.shortcuts import get_object_or_404
@api_view(["GET"])
def retrieve_pteam(request, id):
    print("HELLO RETEIVE it")
    pteam=PTeam.objects.filter(lead=id)
    serializer = PTeamSerializer(pteam,many=True)
    return Response(serializer.data)

def update_pteam(request, pk):
    pteam = get_object_or_404(PTeam, pk=pk)
    if request.method == 'PUT':
        serializer = PTeamSerializer(pteam, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
def delete_pteam(request, pk):
    pteam = get_object_or_404(PTeam, pk=pk)
    if request.method == 'DELETE':
        pteam.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


@api_view(['GET'])
def UserTeam(request,id):
    Userrole=UserRole.objects.get(user=id)
    pteam=PTeam.objects.filter(lead=Userrole.teams.lead.id)
    serializer = PTeamSerializer(pteam,many=True)
    return Response(serializer.data)

@api_view(["POST"])
def selectTask(request):
    task_serializer=TaskSerializer(data=request.data)
    usid=UserStory.objects.get(id=request.data["usid"])
    usid.status=True
    usid.save()
    if task_serializer.is_valid():
        task_serializer.save()
        print("TASK HAS BEEN CREATED for",request.data["assinged_to"])
        return Response(task_serializer.data,status=status.HTTP_201_CREATED)
    print(task_serializer.errors)
    return Response(task_serializer.errors,status=status.HTTP_400_BAD_REQUEST)
    
@api_view(["GET"])
def getTeams(request):
    team=Team.objects.all()
    team=TeamSerializer(team,many=True)
    return Response(team.data,status=status.HTTP_200_OK)

@api_view(["GET"])
def getDevelopers(request,id):
    team =Team.objects.get(lead=id)
    users=UserRole.objects.filter(teams=team,role="developer")
    users=RoleSerializer(users,many=True)
    return Response(users.data,status=status.HTTP_200_OK)

import requests
def download_text_file(file_id, save_path):
    url = f"https://drive.google.com/uc?id={file_id}"
    response = requests.get(url)
    if response.status_code == 200:
        with open(save_path, 'wb') as f:
            f.write(response.content)
        print("File downloaded successfully.")
    else:
        print("Failed to download the file.")

def read_text_file(file_path):
    try:
        with open(file_path, 'r') as file:
            content = file.read()
        return content
    except FileNotFoundError:
        print("File not found.")
        return None
    
@api_view(["GET"])
def getUrl(request):
    file_id = '18hIdHYRnniaSwu5KTWiXfA8pb6Pm9KI5'
    save_path = 'file.txt'
    download_text_file(file_id, save_path)
    file_path = 'file.txt'  # Path to the downloaded text file
    file_content = read_text_file(file_path)
    return Response({
        "api":file_content}
        )
@api_view(["GET"])
def search_view(request,query):
    print("HELLLOOOO______--------")
    results = UserRole.objects.filter(user__username__icontains=query)
    user=RoleSerializer(results,many=True)
    return Response(user.data,status=status.HTTP_200_OK)